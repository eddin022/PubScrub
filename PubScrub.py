import requests
from bs4 import BeautifulSoup
import openai
from docx import Document

#user input
pubMedPrompt = input('Enter your PubMed search prompt: ')
#openai.api_key = input('Enter your openai.api_key: ')
openai.api_key = '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'


#get pubmed URL
pubMedPrompt = pubMedPrompt.replace(' ','+')
url = 'https://pubmed.ncbi.nlm.nih.gov/?term='+pubMedPrompt+'&sort=pubdate&sort_order=asc&size=200'
print('-Prompt: "'+pubMedPrompt+'" -')

# Send GET request and retrieve the HTML content
response = requests.get(url)
soup = BeautifulSoup( response.content , 'html.parser')

#get all pubmed id's
allPMIDs = soup.findAll('span','docsum-pmid')

#get all title, year, and abstract info
pmIDs = []
for ids in allPMIDs:
    pmIDs.append(ids.get_text().strip())
print('Number of abstracts found:'+str(len(pmIDs)))
abstractText = []
worked = []
allYears=[]
allTitles=[]
i=0
print('-Extracting abstracts-')
for ids in pmIDs:
    url = 'https://pubmed.ncbi.nlm.nih.gov/'+ids+'/'
    response = requests.get(url)
    soup = BeautifulSoup( response.content, 'html.parser')
    if soup.find('div','abstract-content selected'):
        abstractText.append(soup.find('div','abstract-content selected').get_text().strip())
    else:
        abstractText.append('No Abstract Found')
    if soup.find('span','cit'):
        temp = soup.find('span','cit').get_text().strip().split(' ')[0]
        temp = temp.split(';')[0]
        allYears.append(temp)
    else:
        allYears.append('No Year Found')
    if soup.find('h1','heading-title'):
        allTitles.append(soup.find('h1','heading-title').get_text().strip())
    else:
        allTitles.append('No Title Found')
            
    i=i+1
    print(str(i)+'/'+str(len(pmIDs)))
print('-Abstracts extracted-')

#use openai api to ask ChatGPT to summarize the abstracts
summaries=[]
print('-ChatGPT is summarizing '+str(len(pmIDs))+' abstracts-')
j=0
for i in abstractText:
    j=j+1
    print(str(j)+'/'+str(len(pmIDs)))
    response = openai.ChatCompletion.create(
      model = "gpt-3.5-turbo",
      #max_tokens = 1000,
      messages = [
        {'role':'system','content': 'You are a PhD level scientist who needs to provide a complete yet concise summary of the scientific abstracts presented to you. Please refrain from saying the phrase "This study".'},
        {"role": "user", "content": 'Please generate a concise summary of the novel findings in this abstract in a maximum of 3 sentences: '+i}
      ]
    )
    
    summaries.append(response['choices'][0]['message']['content'])
print('-Abstracts summarized-')

#save information to a word document
document = Document()
document.add_heading('Summarized PubMed abstracts using the prompt: '+pubMedPrompt)

for i in range(len(pmIDs)):
    title=allTitles[i]
    year=allYears[i]
    summary=summaries[i]
    document.add_heading(title)
    document.add_paragraph('-'+year+'-')
    document.add_paragraph(summary)
    
document.save('PubMed Abstracts - '+pubMedPrompt+'.docx')
print('-PubScraped!-')




